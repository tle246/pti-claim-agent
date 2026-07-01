import argparse
import os
import re
import shutil
import zipfile
import docx
from PIL import Image
from docx2pdf import convert

def parse_antigravity_md(filepath):
    config = {}
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
        
    # Extract Policy Number
    m = re.search(r'-\s+Policy number:\s+`([^`]+)`', content)
    if m:
        config['POLICY_NUMBER'] = m.group(1)
        
    # Extract Bank Account
    m = re.search(r'-\s+Account:\s+`([^`]+)`', content)
    if m:
        config['BANK_ACCOUNT'] = m.group(1)
        
    # Extract Bank Name
    m = re.search(r'-\s+Bank:\s+`([^`]+)`', content)
    if m:
        config['BANK_NAME'] = m.group(1)
        
    # Extract Bank Beneficiary
    m = re.search(r'-\s+Beneficiary:\s+`([^`]+)`', content)
    if m:
        config['BANK_BENEFICIARY'] = m.group(1)
        
    return config

def fill_claim_form(args):
    workspace = "/Users/dile/Downloads/project_app/pti-claim-agent"
    antigravity_path = os.path.join(workspace, "ANTIGRAVITY.md")
    config = parse_antigravity_md(antigravity_path)
    
    insured_name = args.name
    insured_dob = args.dob
    insured_cccd = args.cccd
    visit_date = args.visit_date
    hospital = args.hospital
    diagnosis = args.diagnosis
    amount_vnd = args.amount
    
    # 1. Create output folder
    folder_name = f"LeVanTinh_2026-06-03"
    output_dir = os.path.join(workspace, "output", folder_name)
    os.makedirs(output_dir, exist_ok=True)
    
    # 2. Copy blank form to output
    doc_path = os.path.join(output_dir, f"GYC_{insured_name.replace(' ', '')}_{visit_date.replace('/', '-')}.docx")
    shutil.copy(os.path.join(workspace, "assets", "blank_form.docx"), doc_path)
    
    # 3. Fill text fields
    doc = docx.Document(doc_path)
    mapping = {
        "{{POLICY_NUMBER}}": config.get("POLICY_NUMBER", ""),
        "{{INSURED_NAME}}": insured_name,
        "{{INSURED_DOB}}": insured_dob,
        "{{INSURED_CCCD}}": insured_cccd,
        "{{AMOUNT_VND}}": amount_vnd,
        "{{BANK_ACCOUNT}}": args.bank_account if args.bank_account else config.get("BANK_ACCOUNT", ""),
        "{{BANK_NAME}}": args.bank_name if args.bank_name else config.get("BANK_NAME", ""),
        "{{BANK_BENEFICIARY}}": args.bank_beneficiary if args.bank_beneficiary else config.get("BANK_BENEFICIARY", ""),
        "{{VISIT_DATE}}": visit_date,
        "{{DIAGNOSIS}}": diagnosis,
        "{{HOSPITAL}}": hospital
    }
    
    # Replace in paragraphs
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
                
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in mapping.items():
                        if k in p.text:
                            p.text = p.text.replace(k, v)
                            
    doc.save(doc_path)
    print("Filled text placeholders successfully.")
    
    # 4. Embed signature using ZIP manipulation
    signature_path = os.path.join(workspace, "assets", "signature.png")
    if os.path.exists(signature_path):
        # Calculate aspect ratio
        with Image.open(signature_path) as img:
            width, height = img.size
        aspect_ratio = height / width
        target_width_inch = 2.0
        target_height_inch = target_width_inch * aspect_ratio
        
        width_emu = int(target_width_inch * 914400)
        height_emu = int(target_height_inch * 914400)
        
        temp_doc_path = doc_path + ".temp"
        with zipfile.ZipFile(doc_path, 'r') as z_in:
            with zipfile.ZipFile(temp_doc_path, 'w', zipfile.ZIP_DEFLATED) as z_out:
                rels_content = z_in.read("word/_rels/document.xml.rels").decode("utf-8")
                rids = [int(x) for x in re.findall(r'Id="rId(\d+)"', rels_content)]
                next_rid_num = max(rids) + 1 if rids else 1
                sig_rid = f"rId{next_rid_num}"
                
                for item in z_in.infolist():
                    if item.filename.endswith('/'):
                        continue
                        
                    if item.filename == "word/document.xml":
                        doc_content = z_in.read(item.filename).decode("utf-8")
                        
                        drawing_xml = f'''<w:r><w:drawing><wp:inline distT="0" distB="0" distL="0" distR="0"><wp:extent cx="{width_emu}" cy="{height_emu}"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:docPr id="999" name="Signature"/><wp:cNvGraphicFramePr><a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/></wp:cNvGraphicFramePr><a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:nvPicPr><pic:cNvPr id="0" name="signature.png"/><pic:cNvPicPr/></pic:nvPicPr><pic:blipFill><a:blip r:embed="{sig_rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r>'''
                        
                        pattern_w_r = r'<w:r(?:[^>]*)>\s*<w:t(?:[^>]*)>{{\s*SIGNATURE\s*}}</w:t>\s*</w:r>'
                        if re.search(pattern_w_r, doc_content):
                            doc_content = re.sub(pattern_w_r, drawing_xml, doc_content)
                        elif "{{SIGNATURE}}" in doc_content:
                            doc_content = doc_content.replace("{{SIGNATURE}}", rf'</w:t></w:r>{drawing_xml}<w:r><w:t xml:space="preserve">')
                        else:
                            doc_content = doc_content.replace("</w:body>", f"<w:p>{drawing_xml}</w:p></w:body>")
                            
                        z_out.writestr(item.filename, doc_content)
                        
                    elif item.filename == "word/_rels/document.xml.rels":
                        new_rel = f'<Relationship Id="{sig_rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/signature.png"/>'
                        content = rels_content.replace('</Relationships>', f'{new_rel}</Relationships>')
                        z_out.writestr(item.filename, content)
                        
                    elif item.filename == "[Content_Types].xml":
                        content = z_in.read(item.filename).decode("utf-8")
                        if 'Extension="png"' not in content:
                            png_ct = '<Default Extension="png" ContentType="image/png"/>'
                            content = content.replace('</Types>', f'{png_ct}</Types>')
                        z_out.writestr(item.filename, content)
                        
                    else:
                        z_out.writestr(item.filename, z_in.read(item.filename))
                
                # Add signature file
                with open(signature_path, 'rb') as f:
                    sig_bytes = f.read()
                z_out.writestr("word/media/signature.png", sig_bytes)
                
        os.remove(doc_path)
        shutil.move(temp_doc_path, doc_path)
        print("Embedded signature successfully.")
        
        # Convert to PDF
        print("Converting to PDF...")
        pdf_path = doc_path.replace('.docx', '.pdf')
        convert(doc_path, pdf_path)
        if os.path.exists(pdf_path):
            os.remove(doc_path)
            print(f"Generated PDF at {pdf_path}")
        else:
            print("Failed to convert to PDF.")
    else:
        print("Signature file not found at assets/signature.png.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Fill PTI Claim Form")
    parser.add_argument("--name", required=True, help="Insured Name")
    parser.add_argument("--dob", required=True, help="Insured DOB")
    parser.add_argument("--cccd", default="", help="Insured CCCD")
    parser.add_argument("--visit_date", required=True, help="Visit Date")
    parser.add_argument("--hospital", required=True, help="Hospital")
    parser.add_argument("--diagnosis", required=True, help="Diagnosis")
    parser.add_argument("--amount", required=True, help="Amount VND")
    parser.add_argument("--bank_account", default="", help="Bank Account")
    parser.add_argument("--bank_name", default="", help="Bank Name")
    parser.add_argument("--bank_beneficiary", default="", help="Bank Beneficiary")
    args = parser.parse_args()
    fill_claim_form(args)
