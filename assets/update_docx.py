import sys
import docx

def update_docx(input_path, output_path):
    doc = docx.Document(input_path)
    
    mapping = {
        "Hợp đồng bảo hiểm số (Policy No):": "Hợp đồng bảo hiểm số (Policy No): {{POLICY_NUMBER}}",
        "Tên người được bảo hiểm (Name of Insured):": "Tên người được bảo hiểm (Name of Insured): {{INSURED_NAME}}",
        "CCCD/Hộ Chiếu (ID/Passport of Insured):": "CCCD/Hộ Chiếu (ID/Passport of Insured): {{INSURED_CCCD}}",
        "Ngày sinh (D.O.B of Insured):": "Ngày sinh (D.O.B of Insured): {{INSURED_DOB}}",
        "Số tiền yêu cầu chi trả": "Số tiền yêu cầu chi trả (Total amount claimed): {{AMOUNT_VND}}",
        "Account No/ (Số tài khoản):": "Account No/ (Số tài khoản): {{BANK_ACCOUNT}}",
        "Bank name/ (Tên Ngân hàng):": "Bank name/ (Tên Ngân hàng): {{BANK_NAME}}",
        "Beneficiary/ (Người thụ hưởng):": "Beneficiary/ (Người thụ hưởng): {{BANK_BENEFICIARY}}",
        "Ngày nhập viện (Date of admission):": "Ngày nhập viện (Date of admission): {{VISIT_DATE}}",
        "Chẩn đoán bệnh / Nguyên nhân tai nạn (Medical conditions or Diagnosis /Cause of accident):": "Chẩn đoán bệnh / Nguyên nhân tai nạn (Medical conditions or Diagnosis /Cause of accident): {{DIAGNOSIS}}",
        "Tên cơ sở y tế (Name of Hospital or clinic):": "Tên cơ sở y tế (Name of Hospital or clinic): {{HOSPITAL}}",
        "(Chữ ký và họ tên của Người được bảo hiểm)": "(Chữ ký và họ tên của Người được bảo hiểm)\n{{SIGNATURE}}"
    }

    for p in doc.paragraphs:
        text = p.text
        for k, v in mapping.items():
            if k in text and "{{" not in text:
                p.text = text.replace(k, v)
                
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    for k, v in mapping.items():
                        if k in text and "{{" not in text:
                            p.text = text.replace(k, v)
                            
    doc.save(output_path)

if __name__ == "__main__":
    update_docx("Giay-Yeu-Cau-Tra-Tien-Bao-Hiem.docx", "blank_form.docx")
